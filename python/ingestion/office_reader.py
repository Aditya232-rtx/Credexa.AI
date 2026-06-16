import docx
import openpyxl
import pandas as pd
import os

def parse_docx(file_path):
    """
    Extracts text, tables, and metadata from a .docx file.
    """
    doc = docx.Document(file_path)
    
    # 1. Core Properties
    cp = doc.core_properties
    metadata = {
        'author': cp.author,
        'created': cp.created.isoformat() if cp.created else None,
        'modified': cp.modified.isoformat() if cp.modified else None,
        'last_modified_by': cp.last_modified_by,
        'revision': cp.revision
    }
    
    # 2. Text Paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # 3. Tables
    tables = []
    for table in doc.tables:
        t_data = []
        for row in table.rows:
            r_data = [cell.text.strip() for cell in row.cells]
            t_data.append(r_data)
        tables.append(t_data)
        
    return {
        'metadata': metadata,
        'text': "\n".join(paragraphs),
        'tables': tables
    }

def parse_xlsx(file_path):
    """
    Extracts cell values, formulas, and workbook properties from an .xlsx file.
    """
    wb = openpyxl.load_workbook(file_path, data_only=False) # Get formulas
    wb_data = openpyxl.load_workbook(file_path, data_only=True) # Get evaluated values
    
    metadata = {
        'creator': wb.properties.creator,
        'last_modified_by': wb.properties.lastModifiedBy,
        'created': wb.properties.created.isoformat() if wb.properties.created else None,
        'modified': wb.properties.modified.isoformat() if wb.properties.modified else None,
        'revision': wb.properties.revision
    }
    
    sheets_data = {}
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_data_only = wb_data[sheet_name]
        
        sheet_info = {
            'hidden': sheet.sheet_state == 'hidden',
            'rows': []
        }
        
        for row_idx, row in enumerate(sheet.iter_rows()):
            row_data = []
            for col_idx, cell in enumerate(row):
                val_cell = sheet_data_only.cell(row=row_idx+1, column=col_idx+1)
                
                cell_info = {
                    'coordinate': cell.coordinate,
                    'formula': cell.value if cell.data_type == 'f' else None,
                    'value': val_cell.value,
                    'hyperlink': cell.hyperlink.target if cell.hyperlink else None
                }
                if cell_info['value'] is not None or cell_info['formula'] is not None:
                    row_data.append(cell_info)
            if row_data:
                sheet_info['rows'].append(row_data)
                
        sheets_data[sheet_name] = sheet_info
        
    return {
        'metadata': metadata,
        'sheets': sheets_data
    }

def parse_csv(file_path):
    """
    Parses .csv files into standard text and data representations.
    """
    df = pd.read_csv(file_path)
    
    # Simple table representation
    tables = [df.values.tolist()]
    
    # Column headers
    headers = df.columns.tolist()
    tables[0].insert(0, headers)
    
    text = df.to_string(index=False)
    
    return {
        'metadata': {},
        'text': text,
        'tables': tables
    }

def read_office_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return parse_docx(file_path)
    elif ext == '.xlsx':
        return parse_xlsx(file_path)
    elif ext == '.csv':
        return parse_csv(file_path)
    else:
        raise ValueError(f"Unsupported office file format: {ext}")
